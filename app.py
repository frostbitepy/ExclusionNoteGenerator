import docx


file_path = "plantilla.docx"
xfechax = "xfechax"
xentidadx = "xentidadx"
xreceptorx = "xreceptorx"
xproductox = "xproductox"
xmonedax = "xmonedax"
xmesx = "xmesx"
xnronotax = "xnronotax"
xproductox = "xproductox"
xmonedax = "xmonedax"
xmesx = "xmesx"


document = docx.Document(file_path)


def replace_date(fecha):
    for paragraph in document.paragraphs:
        if xfechax in paragraph.text:
            paragraph.text = paragraph.text.replace(xfechax, fecha)
    #document.save(new_file)

def replace_entidad(entidad):
    #document = docx.Document(file_path)
    for paragraph in document.paragraphs:
        if xentidadx in paragraph.text:
            paragraph.text = paragraph.text.replace(xentidadx, entidad)
    #document.save(new_file)

def replace_receptor(receptor):
    #document = docx.Document(file_path)
    for paragraph in document.paragraphs:
        if xreceptorx in paragraph.text:
            paragraph.text = paragraph.text.replace(xreceptorx, receptor)
    #document.save(new_file)

def replace_nota(nota):
    #document = docx.Document(file_path)
    for paragraph in document.paragraphs:
        if xnronotax in paragraph.text:
            paragraph.text = paragraph.text.replace(xnronotax, nota)
    #document.save(new_file)

def replace_producto(producto):
    #document = docx.Document(file_path)
    for paragraph in document.paragraphs:
        if xproductox in paragraph.text:
            paragraph.text = paragraph.text.replace(xproductox, producto)
    #document.save(new_file)

def replace_moneda(moneda):
    #document = docx.Document(file_path)
    for paragraph in document.paragraphs:
        if xmonedax in paragraph.text:
            paragraph.text = paragraph.text.replace(xmonedax, moneda)
    #document.save(new_file)

def replace_mes(mes):
    #document = docx.Document(file_path)
    for paragraph in document.paragraphs:
        if xmesx in paragraph.text:
            paragraph.text = paragraph.text.replace(xmesx, mes)
    #document.save(new_file)


new_date = "28 de junio de 2023"
new_file = "new_file.docx"
new_entidad = "Sudameris Bank SAECA"
new_receptor = "Alicia Gonzalez"
new_nota = "666"
new_producto = "Cancelación de Deudas"
new_moneda = "Guaranies"
new_mes = "Junio"

replace_date(new_date)
print("The word '{}' has been replaced with '{}' in the file '{}'.".format(xfechax, new_date, new_file))

replace_entidad(new_entidad)
print("The word '{}' has been replaced with '{}' in the file '{}'.".format(xentidadx, new_entidad, new_file))  

replace_receptor(new_receptor)
print("The word '{}' has been replaced with '{}' in the file '{}'.".format(xreceptorx, new_receptor, new_file))

replace_nota(new_nota)
print("The word '{}' has been replaced with '{}' in the file '{}'.".format(xnronotax, new_nota, new_file))

replace_producto(new_producto)
print("The word '{}' has been replaced with '{}' in the file '{}'.".format(xproductox, new_producto, new_file))

replace_moneda(new_moneda)
print("The word '{}' has been replaced with '{}' in the file '{}'.".format(xmonedax, new_moneda, new_file))

replace_mes(new_mes)
print("The word '{}' has been replaced with '{}' in the file '{}'.".format(xmesx, new_mes, new_file))


document.save(new_file)

                  